from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from PIL import Image, ImageDraw
from services.repo_analyzer import RepoAnalyzer

# Official Brand Colors (v2.0)
BRAND_RED = (239, 62, 37)
BRAND_PURPLE = (93, 36, 143)

def set_table_borders(table):
    """Utility to set 1px solid black borders on a table"""
    tbl = table._tbl
    for cell in tbl.xpath('.//w:tc'):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # 1/2 pt = 4, 1pt = 8
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def add_red_box_to_image(image_path, box=None):
    """Draw a 2px red box for generic annotations"""
    try:
        with Image.open(image_path) as img:
            draw = ImageDraw.Draw(img)
            width, height = img.size
            if not box:
                box = [width * 0.2, height * 0.2, width * 0.5, height * 0.4]
            draw.rectangle(box, outline="red", width=2)
            img.save(image_path)
    except Exception as e:
        print(f"Failed to annotate image: {e}")

async def generate_dual_manuals(run, steps, mode="branded"):
    """Entry point with mode selection"""
    if mode == "all":
        # Generate both for maximum availability
        await generate_generic_user_manual(run, steps)
        await generate_generic_tech_manual(run, steps)
        await generate_user_manual(run, steps)
        await generate_tech_manual(run, steps)
    elif mode == "generic":
        user_path = await generate_generic_user_manual(run, steps)
        tech_path = await generate_generic_tech_manual(run, steps)
    else:
        user_path = await generate_user_manual(run, steps)
        tech_path = await generate_tech_manual(run, steps)
    return True

async def generate_generic_user_manual(run, steps):
    doc = Document()
    # Sections...
    doc.add_heading("1. DOCUMENT CONTROL", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table)
    
    data = [
        ["TITLE", "GENERIC USER MANUAL"],
        ["VERSION", "1.0.0 (ENG)"],
        ["DATE", datetime.now().strftime("%Y-%m-%d")],
        ["AUDIENCE", "End Users / Engineering Teams"],
        ["PRODUCT", run.url]
    ]
    for i, row in enumerate(data):
        table.cell(i, 0).text = row[0]
        table.cell(i, 1).text = row[1]
        table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\n")
    doc.add_heading("2. PRODUCT DEFINITION", level=1)
    doc.add_paragraph(f"This document provides a strictly professional engineering-grade guide for the application located at {run.url}. It traces core user journeys and functional interfaces discovered through autonomous recursive exploration.")

    doc.add_heading("3. INTERFACE GUIDE", level=1)
    for i, step in enumerate(steps):
        doc.add_heading(f"Module {i+1}: {step.description}", level=2)
        if step.screenshot_path and os.path.exists(step.screenshot_path):
            add_red_box_to_image(step.screenshot_path)
            doc.add_picture(step.screenshot_path, width=Inches(5))
            cap = doc.add_paragraph(f"Figure {i+1}.1 - Interface analysis of {step.description}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("\n")

    doc.add_heading("4. FEATURE BREAKDOWN", level=1)
    ft_table = doc.add_table(rows=1, cols=2)
    ft_table.style = 'Table Grid'
    set_table_borders(ft_table)
    hdr_cells = ft_table.rows[0].cells
    hdr_cells[0].text = 'Component / Action'
    hdr_cells[1].text = 'Functional Responsibility'
    
    for step in steps:
        row_cells = ft_table.add_row().cells
        row_cells[0].text = step.action
        row_cells[1].text = step.description

    output_dir = "storage/docs"
    os.makedirs(output_dir, exist_ok=True)
    path = f"{output_dir}/generic_user_{run.id}.docx"
    doc.save(path)
    return path

async def generate_generic_tech_manual(run, steps):
    doc = Document()
    analyzer = RepoAnalyzer(".") # Root path analysis
    repo_data = analyzer.analyze()

    header = doc.sections[0].header
    header.paragraphs[0].text = "INTERNAL ENGINEERING USE ONLY"
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("SYSTEM ARCHITECTURE & TECHNICAL SPECIFICATION", level=1)
    
    doc.add_heading("1. TECHNOLOGY STACK", level=2)
    ts_table = doc.add_table(rows=1, cols=2)
    ts_table.style = 'Table Grid'
    set_table_borders(ts_table)
    ts_table.rows[0].cells[0].text = "Dependency"
    ts_table.rows[0].cells[1].text = "Version"
    for item in repo_data["tech_stack"]:
        row = ts_table.add_row().cells
        row[0].text = item["item"]
        row[1].text = item["version"]

    doc.add_heading("2. REPOSITORY LAYOUT", level=2)
    p = doc.add_paragraph()
    p.add_run(repo_data["layout"]).font.name = 'Courier New'

    doc.add_heading("3. DATABASE SCHEMA MODELS", level=2)
    for model in repo_data["db"]:
        doc.add_paragraph(f"Model: {model['name']}").bold = True
        db_table = doc.add_table(rows=1, cols=2)
        db_table.style = 'Table Grid'
        set_table_borders(db_table)
        db_table.rows[0].cells[0].text = "Field Name"
        db_table.rows[0].cells[1].text = "Type"
        for field in model["fields"]:
            row = db_table.add_row().cells
            row[0].text = field["name"]
            row[1].text = field["type"]

    doc.add_heading("4. API REFERENCE", level=2)
    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = 'Table Grid'
    set_table_borders(api_table)
    api_table.rows[0].cells[0].text = "Method"
    api_table.rows[0].cells[1].text = "Path"
    api_table.rows[0].cells[2].text = "Mapped File"
    for route in repo_data["api"]:
        row = api_table.add_row().cells
        row[0].text = route["method"]
        row[1].text = route["path"]
        row[2].text = route["file"]

    path = f"storage/docs/generic_tech_{run.id}.docx"
    doc.save(path)
    return path

# --- PREVIOUS BRANDED GENERATORS ---
async def generate_user_manual(run, steps):
    """Visual-heavy, step-by-step branded user guide"""
    doc = Document()
    
    # Header with Branding
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "MyProBuddy Manual AI - Official User Guide"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title Page
    if run.logo_path and os.path.exists(run.logo_path):
        try:
            doc.add_picture(run.logo_path, width=Inches(2))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    title = doc.add_heading("USER MANUAL", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Target System: {run.url}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("1. INTRODUCTION", level=1)
    doc.add_paragraph("This manual provides a visual, step-by-step guide to the core user journeys of your application, autonomously mapped by MyProBuddy Manual AI.")

    doc.add_heading("2. CORE USER JOURNEYS", level=1)
    for i, step in enumerate(steps):
        doc.add_heading(f"Step {step.step_number}: {step.description}", level=2)
        doc.add_paragraph(f"Action: {step.action}").italic = True
        
        if step.screenshot_path and os.path.exists(step.screenshot_path):
            # Annotate with red box for visual focus
            add_red_box_to_image(step.screenshot_path)
            
            # Center picture
            pic = doc.add_picture(step.screenshot_path, width=Inches(5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Figure caption
            caption = doc.add_paragraph(f"Figure {i+1}.1 - {step.description} Interface")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("\n")

    output_dir = "storage/docs"
    os.makedirs(output_dir, exist_ok=True)
    path = f"{output_dir}/user_manual_{run.id}.docx"
    doc.save(path)
    return path

async def generate_tech_manual(run, steps):
    """Logic-heavy branded technical specification"""
    doc = Document()
    analyzer = RepoAnalyzer(".")
    repo_data = analyzer.analyze()

    # Header
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "MyProBuddy Manual AI - Technical Specification"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = doc.add_heading("TECHNICAL SPECIFICATION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"System: {run.url}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. SYSTEM ARCHITECTURE", level=1)
    doc.add_paragraph("The following architecture was autonomously mapped by analyzing the runtime behavior and the associated source code repository.")

    doc.add_heading("2. REPOSITORY MAPPING", level=2)
    layout_p = doc.add_paragraph()
    layout_p.add_run(repo_data["layout"]).font.name = 'Courier New'

    doc.add_heading("3. LOGIC & CODE TRACING", level=1)
    for step in steps:
        doc.add_heading(f"Module: {step.description}", level=2)
        doc.add_paragraph(f"Action Type: {step.action}")
        doc.add_paragraph(f"Logic: {step.ai_reasoning}").italic = True
        
        if step.mapped_code:
            doc.add_paragraph("Mapped Source Code:").bold = True
            code_p = doc.add_paragraph()
            code_run = code_p.add_run(step.mapped_code)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(8)

    doc.add_heading("4. API ARCHITECTURE", level=1)
    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = 'Table Grid'
    hdr_cells = api_table.rows[0].cells
    hdr_cells[0].text = "Method"
    hdr_cells[1].text = "Path"
    hdr_cells[2].text = "Source File"
    
    for route in repo_data["api"]:
        row = api_table.add_row().cells
        row[0].text = route["method"]
        row[1].text = route["path"]
        row[2].text = route["file"]

    output_dir = "storage/docs"
    os.makedirs(output_dir, exist_ok=True)
    path = f"{output_dir}/tech_manual_{run.id}.docx"
    doc.save(path)
    return path
